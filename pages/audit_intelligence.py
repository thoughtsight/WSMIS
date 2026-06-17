"""
pages.audit_intelligence — Audit Intelligence Report page (PR-027)

One-click AI-powered audit report generation with Markdown/PDF/DOCX export.
Uses the services.ai infrastructure (context → prompt → LLM → Markdown).
"""

import streamlit as st
import pandas as pd
import time
from io import BytesIO
from datetime import datetime

from services.ai import build_context, generate_report, AVAILABLE_PROVIDERS
from utils.filters import apply_month_filter
from utils.constants import ADV_COL
from ui.formatters import fmt_inr


def _render_markdown_download(markdown: str, filename: str = "audit_intelligence_report.md"):
    """Download Markdown file."""
    st.download_button(
        label="📄 Download Markdown",
        data=markdown,
        file_name=filename,
        mime="text/markdown",
        key="download_md",
    )


def _render_pdf_download(markdown: str, filename: str = "audit_intelligence_report.pdf"):
    """Download PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        story = []

        # Custom style for body text
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=12,
        )

        # Custom style for headings
        heading_style = ParagraphStyle(
            "Heading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            spaceAfter=12,
        )

        # Convert Markdown to simple paragraphs (basic parsing)
        lines = markdown.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue
            if line.startswith("# "):
                text = line[2:].strip()
                story.append(Paragraph(text, heading_style))
            elif line.startswith("## "):
                text = line[3:].strip()
                story.append(Paragraph(text, heading_style))
            elif line.startswith("### "):
                text = line[4:].strip()
                story.append(Paragraph(text, heading_style))
            else:
                # Escape special characters for ReportLab
                text = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(text, body_style))

        doc.build(story)
        buffer.seek(0)
        st.download_button(
            label="📕 Download PDF",
            data=buffer,
            file_name=filename,
            mime="application/pdf",
            key="download_pdf",
        )
    except Exception as e:
        st.error(f"PDF generation failed: {e}")


def _render_docx_download(markdown: str, filename: str = "audit_intelligence_report.docx"):
    """Download DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading("Audit Intelligence Report", 0)

        lines = markdown.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.startswith("|") and "|" in line[1:]:
                # Simple table parsing (basic)
                parts = [p.strip() for p in line.split("|") if p.strip()]
                if parts:
                    row = doc.add_table(rows=1, cols=len(parts))
                    row_cells = row.cells
                    for i, part in enumerate(parts):
                        row_cells[i].text = part
            else:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(10)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📘 Download DOCX",
            data=buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx",
        )
    except Exception as e:
        st.error(f"DOCX generation failed: {e}")


def render(df, pairs, alerts, comparison_mode=True, selected_months=None):
    st.markdown('<div class="section-card"><div class="section-title">🧠 Audit Intelligence</div>', unsafe_allow_html=True)

    if df is None or df.empty:
        st.warning("No data available for this period.")
        return

    # Period labels for context
    cp_months = [p[0] for p in pairs]
    pp_months = [p[1] for p in pairs]
    period_label = f"{cp_months[0]} → {cp_months[-1]}" if cp_months else "Current Period"
    comparison_label = "YoY" if comparison_mode else "MoM"

    # Session state for caching
    if "ai_report" not in st.session_state:
        st.session_state.ai_report = None
    if "ai_generation_time" not in st.session_state:
        st.session_state.ai_generation_time = None

    # Generate button
    col1, col2 = st.columns([1, 2])
    with col1:
        generate_btn = st.button("Generate Audit Intelligence Report", key="generate_ai_report", use_container_width=True)

    with col2:
        provider = st.selectbox("AI Provider", AVAILABLE_PROVIDERS, index=0, key="ai_provider")

    if generate_btn:
        with st.spinner("Preparing Audit Intelligence..."):
            start_time = time.time()

            cp = apply_month_filter(df, "Month Name", cp_months)
            pp = apply_month_filter(df, "Month Name", pp_months)

            # Build context and generate report
            try:
                report = generate_report(
                    cp=cp,
                    pp=pp,
                    client_name="Rukmani Motors",
                    period_label=period_label,
                    comparison_label=comparison_label,
                    adv_col=ADV_COL,
                    provider=provider,
                )
                st.session_state.ai_report = report
                st.session_state.ai_generation_time = time.time() - start_time
            except Exception as e:
                st.error(f"Report generation failed: {e}")
                return

    # Display cached report
    if st.session_state.ai_report:
        report = st.session_state.ai_report
        gen_time = st.session_state.ai_generation_time

        # Generation metadata
        st.markdown(f"""
        <div style="background:#F2F2F7;padding:12px;border-radius:8px;margin-bottom:16px;font-size:13px;">
            <strong>Generated in {gen_time:.2f}s</strong> using <strong>{report.provider}</strong> ({report.model})
        </div>
        """, unsafe_allow_html=True)

        # Fallback indicator
        if "offline" in report.provider.lower():
            st.info("🔄 AI provider unavailable — showing offline deterministic report.")

        # Markdown display
        st.markdown(report.markdown)

        # Download buttons
        st.markdown("---")
        st.markdown("### Download Report")
        dl_col1, dl_col2, dl_col3 = st.columns(3)
        with dl_col1:
            _render_markdown_download(report.markdown)
        with dl_col2:
            _render_pdf_download(report.markdown)
        with dl_col3:
            _render_docx_download(report.markdown)

    else:
        st.info("Click 'Generate Audit Intelligence Report' to create a new report.")
